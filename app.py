import streamlit as st
from langchain.prompts import PromptTemplate
from langchain.llms import CTransformers
from docx import Document
from docx.shared import Inches
import io
import requests

# Function to get response from LLama 2 model
def getLLamaresponse(input_text, no_words, blog_style):
    # LLama2 model
    llm = CTransformers(model='C:/Users/GADDA BHAVYASHREE/Downloads/model/llama-2-7b-chat.ggmlv3.q2_K.bin',
                        model_type='llama',
                        config={'max_new_tokens': 256,
                                'temperature': 0.01})

    # Prompt Template
    template = """
        Write a blog for {blog_style} job profile for a topic {input_text}
        within {no_words} words.
            """

    prompt = PromptTemplate(input_variables=["blog_style", "input_text", 'no_words'],
                            template=template)

    # Generate the response from the LLama 2 model
    response = llm(prompt.format(blog_style=blog_style, input_text=input_text, no_words=no_words))
    return response

# Function to fetch image URL
def get_src_original_url(query):
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': "K9RdcbmPcuvUFGHU7hSaWNdwYewuJjwgerRtzV8DxmvOhIdzshVQtuAz",
    }

    params = {
        'query': query,
        'per_page': 1,
    }

    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful (status code 200)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            st.write("No photos found for the given query.")
    else:
        st.write(f"Error: {response.status_code}, {response.text}")

    return None

# Function to create Word document
def create_word_docx(input_text, generated_text, image_url):
    # Create a new Word document
    doc = Document()

    # Add the generated text to the document
    doc.add_heading(input_text, level=1)
    doc.add_paragraph(generated_text)

    # Add the image to the document
    doc.add_picture(io.BytesIO(requests.get(image_url).content), width=Inches(4))

    return doc

st.set_page_config(layout="wide")

def main():
    st.title("AutoBlogger using Llama 2")

    input_text = st.text_input("Enter the Blog Topic")
    no_words = st.text_input('No of Words')
    blog_style = st.selectbox('Writing the blog for', ('Researchers', 'Data Scientist', 'Common People'), index=0)

    submit = st.button("Generate")

    if submit:
        generated_text = getLLamaresponse(input_text, no_words, blog_style)
        image_url = get_src_original_url(input_text)
        if image_url:
            st.image(image_url, caption='Fetched Image', width=500, output_format='PNG') 
        if generated_text:
            st.subheader("Generated Text:")
            st.write(generated_text)
    # Remove the download button

if __name__ == "__main__":
    main()
